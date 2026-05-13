"""Tests for text extractors."""

from __future__ import annotations

from pathlib import Path

import pytest


class TestPDFExtractor:
    """Tests for the PDF extractor."""

    def test_typed_pdf_extraction(self, sample_pdf: Path) -> None:
        """A typed PDF should return non-empty text and success=True."""
        from src.extractor.pdf import PDFExtractor

        result = PDFExtractor().extract(str(sample_pdf))
        assert result["success"] is True
        assert len(result["text"].strip()) > 0

    def test_corrupted_pdf_returns_failure(self, tmp_path: Path) -> None:
        """A corrupted PDF must return success=False and not raise."""
        bad_pdf = tmp_path / "bad.pdf"
        bad_pdf.write_bytes(b"this is not a pdf at all %%%")
        from src.extractor.pdf import PDFExtractor

        result = PDFExtractor().extract(str(bad_pdf))
        assert result["success"] is False
        assert result["error"] is not None
        assert result["text"] == ""

    def test_metadata_extracted(self, sample_pdf: Path) -> None:
        """PDF extractor should always return a metadata dict."""
        from src.extractor.pdf import PDFExtractor

        result = PDFExtractor().extract(str(sample_pdf))
        assert isinstance(result["metadata"], dict)
        assert "page_count" in result["metadata"]


class TestImageExtractor:
    """Tests for the image OCR extractor."""

    def test_image_extraction_returns_result(self, sample_image: Path) -> None:
        """Image extractor should return a result dict without raising."""
        from src.extractor.image import ImageExtractor

        result = ImageExtractor().extract(str(sample_image))
        # Either success (Tesseract installed) or graceful failure
        assert isinstance(result["success"], bool)
        assert isinstance(result["text"], str)

    def test_corrupted_image_returns_failure(self, tmp_path: Path) -> None:
        """A corrupted image must return success=False and not raise."""
        bad_img = tmp_path / "bad.png"
        bad_img.write_bytes(b"not an image")
        from src.extractor.image import ImageExtractor

        result = ImageExtractor().extract(str(bad_img))
        assert result["success"] is False
        assert result["text"] == ""


class TestDocxExtractor:
    """Tests for the DOCX extractor."""

    def test_docx_extraction(self, tmp_path: Path) -> None:
        """A valid DOCX file should be extracted successfully."""
        pytest.importorskip("docx")
        from docx import Document
        from src.extractor.docx import DocxExtractor

        doc = Document()
        doc.add_paragraph("Salary slip for employee John Doe")
        doc.add_paragraph("Net Pay: Rs. 50,000")
        path = tmp_path / "salary.docx"
        doc.save(str(path))

        result = DocxExtractor().extract(str(path))
        assert result["success"] is True
        assert "salary" in result["text"].lower() or "john" in result["text"].lower()

    def test_corrupted_docx_returns_failure(self, tmp_path: Path) -> None:
        """A corrupted DOCX must return success=False and not raise."""
        bad_docx = tmp_path / "bad.docx"
        bad_docx.write_bytes(b"PK\x03\x04bad content")
        from src.extractor.docx import DocxExtractor

        result = DocxExtractor().extract(str(bad_docx))
        assert result["success"] is False


class TestTextExtractor:
    """Tests for the plain text extractor."""

    def test_plain_text(self, tmp_path: Path) -> None:
        """A plain text file should be read correctly."""
        txt = tmp_path / "note.txt"
        txt.write_text("Hello world. This is a test file.", encoding="utf-8")
        from src.extractor.text import TextExtractor

        result = TextExtractor().extract(str(txt))
        assert result["success"] is True
        assert "hello world" in result["text"].lower()

    def test_whatsapp_export_stripped(self, tmp_path: Path) -> None:
        """WhatsApp exports should have timestamps stripped."""
        chat = tmp_path / "chat.txt"
        chat.write_text(
            "01/01/2024, 10:00 - Alice: Hello how are you?\n"
            "01/01/2024, 10:01 - Bob: Fine thanks!\n"
            "01/01/2024, 10:02 - Alice: Great to hear that.\n"
            "01/01/2024, 10:03 - Bob: See you tomorrow.\n"
            "01/01/2024, 10:04 - Alice: Okay bye!\n",
            encoding="utf-8",
        )
        from src.extractor.text import TextExtractor

        result = TextExtractor().extract(str(chat))
        assert result["success"] is True
        # Timestamps and names should be removed
        assert "01/01/2024" not in result["text"]
        assert "Alice" not in result["text"]
