"""DOCX text extractor using python-docx."""

from __future__ import annotations

from typing import Any

from src.extractor.base import BaseExtractor
from src.utils.helpers import clean_text
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxExtractor(BaseExtractor):
    """Extract text from .docx files including tables."""

    def extract(self, file_path: str) -> dict[str, Any]:
        """Extract paragraphs and table contents from a DOCX file."""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            return self._fail("python-docx is not installed")

        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            return self._fail("Corrupted or invalid DOCX file")
        except Exception as exc:
            return self._fail(f"Cannot open DOCX: {exc}")

        try:
            parts: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            text = "\n".join(parts)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            return self._ok(clean_text(text), metadata)

        except Exception as exc:
            logger.warning("Error extracting DOCX %s: %s", file_path, exc)
            return self._fail(str(exc))
